from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import JSONResponse
import pandas as pd
import numpy as np
import json
import io
from docx import Document
from docx.shared import Pt

app = FastAPI()

import os

base_dir = os.path.dirname(os.path.abspath(__file__))
param_path = os.path.join(base_dir, "parameters.json")

with open(param_path, "r", encoding="utf-8") as f:
    parameters = json.load(f)

def convert_coordinates(X, Y, Z, dX, dY, dZ, wx, wy, wz, m, to_gsk):
    if not to_gsk:
        m = -m
        wx, wy, wz = -wx, -wy, -wz
        dX, dY, dZ = -dX, -dY, -dZ

    R = np.array([
        [1, wz, -wy],
        [-wz, 1, wx],
        [wy, -wx, 1]
    ])

    input_coords = np.array([X, Y, Z])
    transformed = (1 + m) * R @ input_coords + np.array([dX, dY, dZ])
    return transformed[0], transformed[1], transformed[2]

def create_docx_report(from_system, to_system, result_df):
    doc = Document()
    doc.add_heading('Результат преобразования координат', level=1)
    doc.add_paragraph(f'Исходная система: {from_system}')
    doc.add_paragraph(f'Целевая система: {to_system}')
    doc.add_paragraph('Первые 5 строк результата:')

    table = doc.add_table(rows=1, cols=len(result_df.columns))
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(result_df.columns):
        hdr_cells[i].text = col_name

    for _, row in result_df.head().iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f.read()

@app.post("/convert")
async def convert(
    file: UploadFile = File(...),
    from_system: str = "СК-42",
    to_system: str = "ГСК-2011"
):
    if not file.filename.endswith((".xlsx", ".xls")):
        raise HTTPException(status_code=400, detail="Требуется файл Excel (.xlsx или .xls)")

    try:
        contents = await file.read()
        df = pd.read_excel(io.BytesIO(contents))

        required_columns = ['X', 'Y', 'Z']
        if not all(col in df.columns for col in required_columns):
            raise HTTPException(status_code=400, detail=f"Файл должен содержать колонки: {required_columns}")

        converted = []

        for _, row in df.iterrows():
            X, Y, Z = row['X'], row['Y'], row['Z']

            if to_system == "ГСК-2011":
                p = parameters[from_system]
                res = convert_coordinates(X, Y, Z,
                                          p["dX"], p["dY"], p["dZ"],
                                          np.radians(p["wx"] / 3600),
                                          np.radians(p["wy"] / 3600),
                                          np.radians(p["wz"] / 3600),
                                          p["m"],
                                          to_gsk=True)
            elif from_system == "ГСК-2011":
                p = parameters[to_system]
                res = convert_coordinates(X, Y, Z,
                                          p["dX"], p["dY"], p["dZ"],
                                          np.radians(p["wx"] / 3600),
                                          np.radians(p["wy"] / 3600),
                                          np.radians(p["wz"] / 3600),
                                          p["m"],
                                          to_gsk=False)
            else:
                p_from = parameters[from_system]
                X1, Y1, Z1 = convert_coordinates(X, Y, Z,
                                                 p_from["dX"], p_from["dY"], p_from["dZ"],
                                                 np.radians(p_from["wx"] / 3600),
                                                 np.radians(p_from["wy"] / 3600),
                                                 np.radians(p_from["wz"] / 3600),
                                                 p_from["m"],
                                                 to_gsk=True)

                p_to = parameters[to_system]
                res = convert_coordinates(X1, Y1, Z1,
                                          p_to["dX"], p_to["dY"], p_to["dZ"],
                                          np.radians(p_to["wx"] / 3600),
                                          np.radians(p_to["wy"] / 3600),
                                          np.radians(p_to["wz"] / 3600),
                                          p_to["m"],
                                          to_gsk=False)

            converted.append(res)

        result_df = pd.DataFrame(converted, columns=["X", "Y", "Z"])

        stream = io.StringIO()
        result_df.to_csv(stream, index=False)

        report_md = f"""## 📊 Результат преобразования

### Исходная система: `{from_system}`
### Целевая система: `{to_system}`

#### Первые 5 строк результата:
{result_df.head().to_markdown(index=False)}"""

        docx_bytes = create_docx_report(from_system, to_system, result_df)

        # Отправляем docx в hex-строке
        return JSONResponse(content={
            "csv": stream.getvalue(),
            "report": report_md,
            "docx": docx_bytes.hex()
        })

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
